from typing import Annotated
from fastmcp import FastMCP
from fastmcp.server.auth.providers.bearer import BearerAuthProvider, RSAKeyPair
import markdownify
from mcp import ErrorData, McpError
from mcp.server.auth.provider import AccessToken
from mcp.types import INTERNAL_ERROR, INVALID_PARAMS, TextContent
from openai import BaseModel
from pydantic import AnyUrl, Field
import readabilipy
from pathlib import Path
import os
import re

TOKEN = "4b1b17db54e3"
MY_NUMBER = "919182243523"  # Insert your number {91}{Your number}


class RichToolDescription(BaseModel):
    description: str
    use_when: str
    side_effects: str | None


class SimpleBearerAuthProvider(BearerAuthProvider):
    """
    A simple BearerAuthProvider that does not require any specific configuration.
    It allows any valid bearer token to access the MCP server.
    For a more complete implementation that can authenticate dynamically generated tokens,
    please use `BearerAuthProvider` with your public key or JWKS URI.
    """

    def __init__(self, token: str):
        k = RSAKeyPair.generate()
        super().__init__(
            public_key=k.public_key, jwks_uri=None, issuer=None, audience=None
        )
        self.token = token

    async def load_access_token(self, token: str) -> AccessToken | None:
        if token == self.token:
            return AccessToken(
                token=token,
                client_id="unknown",
                scopes=[],
                expires_at=None,  # No expiration for simplicity
            )
        return None


class ResumeParser:
    """Handles resume file detection, parsing, and conversion to markdown."""
    
    SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.txt', '.md']
    
    @classmethod
    def find_resume_file(cls, search_directory: str = ".") -> Path | None:
        """
        Find resume file in the specified directory.
        Looks for files with 'resume' in the name or common resume filenames.
        """
        search_path = Path(search_directory)
        
        # Common resume filename patterns
        resume_patterns = [
            'resume.*',
            'cv.*',
            'my_resume.*',
            'my_cv.*',
            '*_resume.*',
            '*_cv.*'
        ]
        
        for pattern in resume_patterns:
            for file_path in search_path.glob(pattern):
                if file_path.suffix.lower() in cls.SUPPORTED_EXTENSIONS:
                    return file_path
        
        # If no pattern matches, look for any supported file in the directory
        for ext in cls.SUPPORTED_EXTENSIONS:
            for file_path in search_path.glob(f"*{ext}"):
                return file_path
                
        return None
    
    @classmethod
    async def extract_text_from_pdf(cls, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            import pymupdf  # PyMuPDF
            
            doc = pymupdf.open(file_path)
            text_content = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                if text.strip():
                    text_content.append(text.strip())
            
            doc.close()
            return "\n\n".join(text_content)
            
        except ImportError:
            # Fallback to pdfplumber if pymupdf is not available
            try:
                import pdfplumber
                
                text_content = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text and text.strip():
                            text_content.append(text.strip())
                
                return "\n\n".join(text_content)
                
            except ImportError:
                raise McpError(
                    ErrorData(
                        code=INTERNAL_ERROR,
                        message="PDF parsing libraries not available. Please install pymupdf or pdfplumber."
                    )
                )
        except Exception as e:
            raise McpError(
                ErrorData(
                    code=INTERNAL_ERROR,
                    message=f"Failed to extract text from PDF: {e!r}"
                )
            )
    
    @classmethod
    async def extract_text_from_docx(cls, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
            
        except ImportError:
            raise McpError(
                ErrorData(
                    code=INTERNAL_ERROR,
                    message="DOCX parsing library not available. Please install python-docx."
                )
            )
        except Exception as e:
            raise McpError(
                ErrorData(
                    code=INTERNAL_ERROR,
                    message=f"Failed to extract text from DOCX: {e!r}"
                )
            )
    
    @classmethod
    async def extract_text_from_text_file(cls, file_path: Path) -> str:
        """Extract text from plain text or markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                raise McpError(
                    ErrorData(
                        code=INTERNAL_ERROR,
                        message=f"Failed to read text file: {e!r}"
                    )
                )
        except Exception as e:
            raise McpError(
                ErrorData(
                    code=INTERNAL_ERROR,
                    message=f"Failed to read text file: {e!r}"
                )
            )
    
    @classmethod
    def clean_and_format_text(cls, text: str) -> str:
        """Clean and format extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove trailing spaces from lines
        lines = [line.rstrip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text.strip()
    
    @classmethod
    def convert_to_markdown(cls, text: str, file_extension: str) -> str:
        """Convert text content to markdown format."""
        if not text:
            return "# Resume\n\n*No content could be extracted from the resume file.*"
        
        # If already markdown, return as-is (with minimal cleaning)
        if file_extension.lower() == '.md':
            return cls.clean_and_format_text(text)
        
        # For other formats, apply basic markdown formatting
        lines = text.split('\n')
        markdown_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                markdown_lines.append("")
                continue
            
            # Detect potential headings (all caps, short lines, etc.)
            if (len(line) < 50 and 
                (line.isupper() or 
                 any(keyword in line.lower() for keyword in ['experience', 'education', 'skills', 'contact', 'summary', 'objective', 'projects']))):
                markdown_lines.append(f"## {line}")
            # Detect bullet points
            elif line.startswith(('•', '-', '*', '◦')) or re.match(r'^\s*[•\-\*]\s+', line):
                cleaned_line = re.sub(r'^\s*[•\-\*]\s*', '- ', line)
                markdown_lines.append(cleaned_line)
            else:
                markdown_lines.append(line)
        
        result = '\n'.join(markdown_lines)
        return cls.clean_and_format_text(result)
    
    @classmethod
    async def parse_resume(cls, search_directory: str = ".") -> str:
        """
        Main method to find and parse resume file.
        Returns the resume content as markdown text.
        """
        # Step 1: Find resume file
        resume_file = cls.find_resume_file(search_directory)
        
        if not resume_file:
            raise McpError(
                ErrorData(
                    code=INVALID_PARAMS,
                    message=f"No resume file found in {search_directory}. Supported formats: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
                )
            )
        
        file_extension = resume_file.suffix.lower()
        
        # Step 2: Extract text based on file type
        try:
            if file_extension == '.pdf':
                raw_text = await cls.extract_text_from_pdf(resume_file)
            elif file_extension == '.docx':
                raw_text = await cls.extract_text_from_docx(resume_file)
            elif file_extension in ['.txt', '.md']:
                raw_text = await cls.extract_text_from_text_file(resume_file)
            else:
                raise McpError(
                    ErrorData(
                        code=INVALID_PARAMS,
                        message=f"Unsupported file format: {file_extension}"
                    )
                )
            
            # Step 3: Convert to markdown
            markdown_content = cls.convert_to_markdown(raw_text, file_extension)
            
            if not markdown_content.strip():
                raise McpError(
                    ErrorData(
                        code=INTERNAL_ERROR,
                        message="No readable content could be extracted from the resume file."
                    )
                )
            
            return markdown_content
            
        except McpError:
            # Re-raise MCP errors as-is
            raise
        except Exception as e:
            raise McpError(
                ErrorData(
                    code=INTERNAL_ERROR,
                    message=f"Unexpected error while parsing resume: {e!r}"
                )
            )


class Fetch:
    IGNORE_ROBOTS_TXT = True
    USER_AGENT = "Puch/1.0 (Autonomous)"

    @classmethod
    async def fetch_url(
        cls,
        url: str,
        user_agent: str,
        force_raw: bool = False,
    ) -> tuple[str, str]:
        """
        Fetch the URL and return the content in a form ready for the LLM, as well as a prefix string with status information.
        """
        from httpx import AsyncClient, HTTPError

        async with AsyncClient() as client:
            try:
                response = await client.get(
                    url,
                    follow_redirects=True,
                    headers={"User-Agent": user_agent},
                    timeout=30,
                )
            except HTTPError as e:
                raise McpError(
                    ErrorData(
                        code=INTERNAL_ERROR, message=f"Failed to fetch {url}: {e!r}"
                    )
                )
            if response.status_code >= 400:
                raise McpError(
                    ErrorData(
                        code=INTERNAL_ERROR,
                        message=f"Failed to fetch {url} - status code {response.status_code}",
                    )
                )

            page_raw = response.text

        content_type = response.headers.get("content-type", "")
        is_page_html = (
            "<html" in page_raw[:100] or "text/html" in content_type or not content_type
        )

        if is_page_html and not force_raw:
            return cls.extract_content_from_html(page_raw), ""

        return (
            page_raw,
            f"Content type {content_type} cannot be simplified to markdown, but here is the raw content:\n",
        )

    @staticmethod
    def extract_content_from_html(html: str) -> str:
        """Extract and convert HTML content to Markdown format.

        Args:
            html: Raw HTML content to process

        Returns:
            Simplified markdown version of the content
        """
        ret = readabilipy.simple_json.simple_json_from_html_string(
            html, use_readability=True
        )
        if not ret["content"]:
            return "<error>Page failed to be simplified from HTML</error>"
        content = markdownify.markdownify(
            ret["content"],
            heading_style=markdownify.ATX,
        )
        return content


mcp = FastMCP(
    "My MCP Server",
    auth=SimpleBearerAuthProvider(TOKEN),
)

ResumeToolDescription = RichToolDescription(
    description="Serve your resume in plain markdown format.",
    use_when="Puch (or anyone) asks for your resume; this must return raw markdown, no extra formatting.",
    side_effects=None,
)

@mcp.tool(description=ResumeToolDescription.model_dump_json())
async def resume() -> str:
    """
    Return your resume exactly as markdown text.
    
    Finds resume file in the current directory and converts it to markdown format.
    Supports PDF, DOCX, TXT, and MD file formats.
    """
    try:
        markdown_resume = await ResumeParser.parse_resume(".")
        return markdown_resume
    except McpError:
        # Re-raise MCP errors to return proper error responses
        raise
    except Exception as e:
        # Handle any unexpected errors
        raise McpError(
            ErrorData(
                code=INTERNAL_ERROR,
                message=f"Resume parsing failed: {e!r}"
            )
        )


@mcp.tool
async def validate() -> str:
    """
    NOTE: This tool must be present in an MCP server used by puch.
    """
    return MY_NUMBER


FetchToolDescription = RichToolDescription(
    description="Fetch a URL and return its content.",
    use_when="Use this tool when the user provides a URL and asks for its content, or when the user wants to fetch a webpage.",
    side_effects="The user will receive the content of the requested URL in a simplified format, or raw HTML if requested.",
)


@mcp.tool(description=FetchToolDescription.model_dump_json())
async def fetch(
    url: Annotated[AnyUrl, Field(description="URL to fetch")],
    max_length: Annotated[
        int,
        Field(
            default=5000,
            description="Maximum number of characters to return.",
            gt=0,
            lt=1000000,
        ),
    ] = 5000,
    start_index: Annotated[
        int,
        Field(
            default=0,
            description="On return output starting at this character index, useful if a previous fetch was truncated and more context is required.",
            ge=0,
        ),
    ] = 0,
    raw: Annotated[
        bool,
        Field(
            default=False,
            description="Get the actual HTML content if the requested page, without simplification.",
        ),
    ] = False,
) -> list[TextContent]:
    """Fetch a URL and return its content."""
    url_str = str(url).strip()
    if not url:
        raise McpError(ErrorData(code=INVALID_PARAMS, message="URL is required"))

    content, prefix = await Fetch.fetch_url(url_str, Fetch.USER_AGENT, force_raw=raw)
    original_length = len(content)
    if start_index >= original_length:
        content = "<error>No more content available.</error>"
    else:
        truncated_content = content[start_index : start_index + max_length]
        if not truncated_content:
            content = "<error>No more content available.</error>"
        else:
            content = truncated_content
            actual_content_length = len(truncated_content)
            remaining_content = original_length - (start_index + actual_content_length)
            # Only add the prompt to continue fetching if there is still remaining content
            if actual_content_length == max_length and remaining_content > 0:
                next_start = start_index + actual_content_length
                content += f"\n\n<error>Content truncated. Call the fetch tool with a start_index of {next_start} to get more content.</error>"
    return [TextContent(type="text", text=f"{prefix}Contents of {url}:\n{content}")]


async def main():
    await mcp.run_async(
        "streamable-http",
        host="0.0.0.0",
        port=8085,
    )


if __name__ == "__main__":
    import asyncio

    asyncio.run(main())
